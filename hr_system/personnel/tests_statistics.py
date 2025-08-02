from django.test import TestCase
from rest_framework.test import APITestCase
from django.contrib.auth.models import User
from docx import Document
import io
import datetime

from .models import (
    Division,
    Position,
    Employee,
    StaffingUnit,
    SecondmentRequest,
    EmployeeStatusLog,
    EmployeeStatusType,
    DivisionType,
    UserProfile,
    UserRole,
)
from .services import get_division_statistics


class StatisticsCalculationTest(TestCase):
    def setUp(self):
        # --- Create Divisions ---
        self.dep1 = Division.objects.create(name="Department 1", division_type=DivisionType.DEPARTMENT)
        self.man1_dep1 = Division.objects.create(
            name="Management 1.1", division_type=DivisionType.MANAGEMENT, parent_division=self.dep1
        )
        self.dep2 = Division.objects.create(name="Department 2", division_type=DivisionType.DEPARTMENT)

        # --- Create Positions ---
        self.manager_pos = Position.objects.create(name="Manager", level=10)
        self.clerk_pos = Position.objects.create(name="Clerk", level=20)

        # --- Define Staffing Plan for Department 1 ---
        StaffingUnit.objects.create(division=self.dep1, position=self.manager_pos, quantity=2)
        StaffingUnit.objects.create(division=self.man1_dep1, position=self.clerk_pos, quantity=8)

        # --- Create Employees for Department 1 ---
        self.emp_in_lineup = Employee.objects.create(
            full_name="Сотрудник В Строю",
            division=self.man1_dep1,
            position=self.clerk_pos,
            hired_date=datetime.date(2024, 1, 1),
        )
        self.emp_on_leave = Employee.objects.create(
            full_name="Сотрудник В Отпуске",
            division=self.man1_dep1,
            position=self.clerk_pos,
            hired_date=datetime.date(2024, 1, 1),
        )
        EmployeeStatusLog.objects.create(
            employee=self.emp_on_leave,
            status=EmployeeStatusType.ON_LEAVE,
            date_from=datetime.date(2025, 1, 1),
            date_to=datetime.date(2025, 1, 31),
        )
        self.emp_sick = Employee.objects.create(
            full_name="Сотрудник На Больничном",
            division=self.man1_dep1,
            position=self.clerk_pos,
            hired_date=datetime.date(2024, 1, 1),
        )
        EmployeeStatusLog.objects.create(
            employee=self.emp_sick,
            status=EmployeeStatusType.SICK_LEAVE,
            date_from=datetime.date(2025, 1, 1),
            date_to=datetime.date(2025, 1, 31),
        )
        self.emp_trip = Employee.objects.create(
            full_name="Сотрудник В Командировке",
            division=self.man1_dep1,
            position=self.clerk_pos,
            hired_date=datetime.date(2024, 1, 1),
        )
        EmployeeStatusLog.objects.create(
            employee=self.emp_trip,
            status=EmployeeStatusType.BUSINESS_TRIP,
            date_from=datetime.date(2025, 1, 1),
            date_to=datetime.date(2025, 1, 31),
        )
        self.emp_seconded_out = Employee.objects.create(
            full_name="Сотрудник Откомандирован",
            division=self.man1_dep1,
            position=self.clerk_pos,
            hired_date=datetime.date(2024, 1, 1),
        )
        EmployeeStatusLog.objects.create(
            employee=self.emp_seconded_out,
            status=EmployeeStatusType.SECONDED_OUT,
            date_from=datetime.date(2025, 1, 1),
            secondment_division=self.dep2,
        )
        self.emp_in_lineup_2 = Employee.objects.create(
            full_name="Еще Один В Строю",
            division=self.dep1,
            position=self.manager_pos,
            hired_date=datetime.date(2024, 1, 1),
        )

        # --- Create Employees Seconded INTO Department 1 ---
        emp_to_second_in_1 = Employee.objects.create(
            full_name="Прикомандированный 1",
            division=self.dep2,
            position=self.clerk_pos,
            hired_date=datetime.date(2024, 1, 1),
        )
        SecondmentRequest.objects.create(
            employee=emp_to_second_in_1,
            from_division=self.dep2,
            to_division=self.dep1,
            status="APPROVED",
            date_from=datetime.date(2025, 1, 1),
        )

        emp_to_second_in_2 = Employee.objects.create(
            full_name="Прикомандированный 2 В Отпуске",
            division=self.dep2,
            position=self.clerk_pos,
            hired_date=datetime.date(2024, 1, 1),
        )
        SecondmentRequest.objects.create(
            employee=emp_to_second_in_2,
            from_division=self.dep2,
            to_division=self.man1_dep1,
            status="APPROVED",
            date_from=datetime.date(2025, 1, 1),
        )
        EmployeeStatusLog.objects.create(
            employee=emp_to_second_in_2,
            status=EmployeeStatusType.ON_LEAVE,
            date_from=datetime.date(2025, 1, 10),
            date_to=datetime.date(2025, 1, 20),
        )

    def test_division_statistics_calculation(self):
        test_date = datetime.date(2025, 1, 15)
        stats = get_division_statistics(self.dep1, test_date)

        # Top-level assertions
        self.assertEqual(
            stats["total_staffing"],
            10,
            "Total staffing should be sum of all units in division and sub-divisions.",
        )
        self.assertEqual(
            stats["on_list_count"],
            6,
            "On list count should be all employees homed in the division.",
        )
        self.assertEqual(
            stats["vacant_count"],
            4,
            "Vacant should be staffing minus on-list.",
        )
        self.assertEqual(
            stats["seconded_in_count"],
            2,
            "Seconded-in count should be 2.",
        )
        self.assertEqual(
            stats["in_lineup_count"],
            2,
            "In line-up count should be 2.",
        )

        # Status counts for on-list employees
        status_counts = stats["status_counts"]
        self.assertEqual(status_counts.get(EmployeeStatusType.ON_DUTY_SCHEDULED, 0), 2)
        self.assertEqual(status_counts.get(EmployeeStatusType.ON_LEAVE, 0), 1)
        self.assertEqual(status_counts.get(EmployeeStatusType.SICK_LEAVE, 0), 1)
        self.assertEqual(status_counts.get(EmployeeStatusType.BUSINESS_TRIP, 0), 1)
        self.assertEqual(status_counts.get(EmployeeStatusType.SECONDED_OUT, 0), 1)

        # Formula check
        self.assertEqual(stats["on_list_count"], sum(status_counts.values()))

        # Seconded-in breakdown
        seconded_status_counts = stats["seconded_in_status_counts"]
        self.assertEqual(
            seconded_status_counts.get(EmployeeStatusType.ON_DUTY_SCHEDULED, 0), 1
        )
        self.assertEqual(seconded_status_counts.get(EmployeeStatusType.ON_LEAVE, 0), 1)
        self.assertEqual(sum(seconded_status_counts.values()), 2)


class ReportGenerationTest(APITestCase):
    def setUp(self):
        self.dep = Division.objects.create(
            name="Report Test Dep", division_type=DivisionType.DEPARTMENT
        )
        self.user = User.objects.create_user(username="testuser", password="password")
        UserProfile.objects.create(user=self.user, role=UserRole.ROLE_4)
        self.client.force_authenticate(user=self.user)

    def test_report_endpoint_returns_docx_file(self):
        url = f"/api/personnel/divisions/{self.dep.id}/report/"
        response = self.client.get(url)

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response["Content-Type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertIn("attachment; filename=", response["Content-Disposition"])

        doc_buffer = io.BytesIO(response.content)
        document = Document(doc_buffer)

        title_text = document.paragraphs[0].text
        self.assertIn("Report Test Dep", title_text)
        self.assertIn("САПТЫҚ ТІЗІМІ", title_text)

        self.assertTrue(len(document.tables) > 0, "The report should contain a table.")
        table = document.tables[0]
        headers = [cell.text for cell in table.rows[0].cells]
        self.assertIn("Количество по списку", headers)
        self.assertIn("В строю", headers)

        # Checking that a status cell has the expected multiline structure
        # Adjust index as needed depending on column ordering
        status_cell = table.rows[1].cells[6]
        self.assertIn(
            "0\nПодстрока 1\nПодстрока 2\nПодстрока 3\nПодстрока 4",
            status_cell.text,
        )

    def test_periodic_report_endpoint(self):
        date_from = "2025-01-10"
        date_to = "2025-01-12"  # 3 days
        url = f"/api/personnel/divisions/{self.dep.id}/periodic-report/?date_from={date_from}&date_to={date_to}"
        response = self.client.get(url)
        self.assertEqual(response.status_code, 200)

        doc_buffer = io.BytesIO(response.content)
        document = Document(doc_buffer)

        titles = [p.text for p in document.paragraphs if "САПТЫҚ ТІЗІМІ" in p.text]
        self.assertEqual(len(titles), 3)
        self.assertIn("10.01.2025", titles[0])
        self.assertIn("11.01.2025", titles[1])
        self.assertIn("12.01.2025", titles[2])

    def test_report_endpoint_returns_xlsx_file(self):
        url = f"/api/personnel/divisions/{self.dep.id}/report/"
        response = self.client.get(url, {"format": "xlsx"})
        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response["Content-Type"],
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

    def test_report_endpoint_returns_pdf_file(self):
        url = f"/api/personnel/divisions/{self.dep.id}/report/"
        response = self.client.get(url, {"format": "pdf"})
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response["Content-Type"], "application/pdf")
