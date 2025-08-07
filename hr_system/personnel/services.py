"""
Сервисные функции для генерации отчетов и расчета статистики.

Полная реализация всех форматов отчетов согласно ТЗ:
- DOCX с таблицей в альбомной ориентации
- XLSX с форматированием
- PDF отчеты
"""

import io
import datetime
from collections import defaultdict, OrderedDict
from typing import Dict, List, Optional, Tuple

from django.db.models import Count, Q, F
from django.utils import timezone

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, PageBreak
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from .models import (
    Division, Employee, EmployeeStatusLog, StaffingUnit,
    SecondmentRequest, EmployeeStatusType, DivisionType,
    SecondmentStatus
)


def get_division_statistics(division: Division, date: datetime.date) -> Dict:
    """
    Получить статистику по подразделению на указанную дату.

    Возвращает словарь с ключами:
    - total_staffing: общее количество по штату
    - on_list_count: количество по списку
    - vacant_count: количество вакансий
    - in_lineup_count: количество в строю
    - status_counts: словарь с количеством по каждому статусу
    - seconded_in_count: количество прикомандированных
    - seconded_in_status_counts: статусы прикомандированных
    """
    # Получаем все штатные единицы подразделения
    staffing_units = StaffingUnit.objects.filter(division=division)
    total_staffing = sum(unit.quantity for unit in staffing_units)

    # Получаем активных сотрудников подразделения
    employees = Employee.objects.filter(division=division, is_active=True)
    on_list_count = employees.count()

    # Считаем вакансии
    vacant_count = total_staffing - on_list_count

    # Собираем статистику по статусам основных сотрудников
    status_counts = defaultdict(int)
    for employee in employees:
        current_status = employee.get_current_status(date)
        status_counts[current_status] += 1

    # Количество в строю
    in_lineup_count = status_counts.get(EmployeeStatusType.ON_DUTY_SCHEDULED, 0)

    # Получаем прикомандированных сотрудников
    seconded_requests = SecondmentRequest.objects.filter(
        to_division=division,
        status=SecondmentStatus.APPROVED,
        date_from__lte=date
    ).filter(
        Q(date_to__gte=date) | Q(date_to__isnull=True)
    )

    seconded_in_count = seconded_requests.count()
    seconded_in_status_counts = defaultdict(int)

    for request in seconded_requests:
        employee = request.employee
        current_status = employee.get_current_status(date)
        seconded_in_status_counts[current_status] += 1

    return {
        'total_staffing': total_staffing,
        'on_list_count': on_list_count,
        'vacant_count': vacant_count,
        'in_lineup_count': in_lineup_count,
        'status_counts': dict(status_counts),
        'seconded_in_count': seconded_in_count,
        'seconded_in_status_counts': dict(seconded_in_status_counts)
    }


def _collect_division_data(division: Division, date: datetime.date) -> Dict:
    """
    Собрать данные по подразделению для отчета.

    Возвращает словарь с полной информацией для строки отчета.
    """
    stats = get_division_statistics(division, date)

    # Собираем детальную информацию по каждому статусу
    status_details = {}

    # Основные сотрудники
    employees = Employee.objects.filter(division=division, is_active=True)
    for status_type in EmployeeStatusType.values:
        employees_with_status = []
        comments = []
        date_ranges = []

        for employee in employees:
            status_log = employee.status_logs.filter(
                status=status_type,
                date_from__lte=date
            ).filter(
                Q(date_to__gte=date) | Q(date_to__isnull=True)
            ).first()

            if status_log:
                employees_with_status.append(employee.full_name)
                if status_log.comment:
                    comments.append(status_log.comment)
                date_range = f"{status_log.date_from.strftime('%d.%m')}"
                if status_log.date_to:
                    date_range += f"-{status_log.date_to.strftime('%d.%m')}"
                date_ranges.append(date_range)

        status_details[status_type] = {
            'count': len(employees_with_status),
            'names': employees_with_status,
            'comments': comments,
            'dates': date_ranges
        }

    # Прикомандированные
    seconded_employees = []
    seconded_requests = SecondmentRequest.objects.filter(
        to_division=division,
        status=SecondmentStatus.APPROVED,
        date_from__lte=date
    ).filter(
        Q(date_to__gte=date) | Q(date_to__isnull=True)
    )

    for request in seconded_requests:
        seconded_employees.append({
            'name': request.employee.full_name,
            'from_division': request.from_division.name,
            'status': request.employee.get_current_status(date)
        })

    return {
        'division': division,
        'stats': stats,
        'status_details': status_details,
        'seconded_employees': seconded_employees
    }


def _create_report_table_docx(doc: Document, data: List[Dict], report_date: datetime.date):
    """Создать таблицу отчета в документе Word согласно ТЗ п.12"""

    # Заголовок согласно ТЗ п.12.2
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Используем первое подразделение для названия
    division_name = data[0]['division'].name if data else "ОРГАНИЗАЦИЯ"

    run = heading.add_run(
        f"{division_name} ЖЕКЕ ҚҰРАМЫНЫҢ САПТЫҚ ТІЗІМІ "
        f"{report_date.strftime('%d.%m.%Y')} ЖЫЛҒЫ"
    )
    run.font.size = Pt(16)
    run.font.bold = True

    # Колонки согласно ТЗ п.12.3
    status_columns = [
        ('На дежурстве', EmployeeStatusType.ON_DUTY_ACTUAL),
        ('После дежурства', EmployeeStatusType.AFTER_DUTY),
        ('В командировке', EmployeeStatusType.BUSINESS_TRIP),
        ('Учёба/соревнования/конференция', EmployeeStatusType.TRAINING_ETC),
        ('В отпуске', EmployeeStatusType.ON_LEAVE),
        ('На больничном', EmployeeStatusType.SICK_LEAVE),
        ('Прикомандирован', EmployeeStatusType.SECONDED_IN),
        ('Откомандирован', EmployeeStatusType.SECONDED_OUT),
    ]

    # Создаем таблицу с учетом 4 строк данных для каждого статуса
    num_cols = 6 + len(status_columns)
    num_rows = 1  # Заголовок

    # Считаем строки: для каждого подразделения нужно 4 строки
    for div_data in data:
        num_rows += 4  # 4 строки для каждого управления
    num_rows += 1  # Итоговая строка

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Заголовки колонок
    header_cells = table.rows[0].cells
    headers = [
        '№',
        'Название управления',
        'Количество по штату',
        'Количество по списку',
        'Вакантные должности',
        'В строю'
    ]

    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        header_cells[idx].paragraphs[0].runs[0].font.bold = True
        header_cells[idx].paragraphs[0].runs[0].font.size = Pt(12)
        header_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Заголовки статусов
    for idx, (status_name, _) in enumerate(status_columns):
        header_cells[6 + idx].text = status_name
        header_cells[6 + idx].paragraphs[0].runs[0].font.bold = True
        header_cells[6 + idx].paragraphs[0].runs[0].font.size = Pt(12)
        header_cells[6 + idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Данные по подразделениям
    row_idx = 1
    totals = defaultdict(int)

    for div_idx, div_data in enumerate(data, start=1):
        division = div_data['division']
        stats = div_data['stats']
        status_details = div_data['status_details']

        # Объединяем ячейки для первых 6 колонок (они охватывают 4 строки)
        for col_idx in range(6):
            for merge_row in range(1, 4):
                if row_idx + merge_row < len(table.rows):
                    cell_a = table.cell(row_idx, col_idx)
                    cell_b = table.cell(row_idx + merge_row, col_idx)
                    cell_a.merge(cell_b)

        # Заполняем основные данные
        cells = table.rows[row_idx].cells
        cells[0].text = str(div_idx)
        cells[1].text = division.name
        cells[2].text = str(stats['total_staffing'])

        # Количество по списку с прикомандированными
        seconded_in = stats['seconded_in_count']
        cells[3].text = f"{stats['on_list_count']}"
        if seconded_in > 0:
            cells[3].text += f" +{seconded_in}"

        cells[4].text = str(stats['vacant_count'])
        cells[5].text = str(stats['in_lineup_count'])

        # Обновляем итоги
        totals['staffing'] += stats['total_staffing']
        totals['on_list'] += stats['on_list_count']
        totals['seconded_in'] += stats['seconded_in_count']
        totals['vacant'] += stats['vacant_count']
        totals['in_lineup'] += stats['in_lineup_count']

        # Данные по статусам (4 строки для каждого)
        for col_idx, (_, status_type) in enumerate(status_columns):
            details = status_details.get(status_type, {})

            # Специальная обработка для прикомандированных
            if status_type == EmployeeStatusType.SECONDED_IN:
                count = stats['seconded_in_count']
                names = [emp['name'] for emp in div_data.get('seconded_employees', [])]
                comments = [f"Из {emp['from_division']}" for emp in div_data.get('seconded_employees', [])]
                dates = []
            else:
                count = details.get('count', 0)
                names = details.get('names', [])
                comments = details.get('comments', [])
                dates = details.get('dates', [])

            totals[f'status_{status_type}'] += count

            # Заполняем 4 строки для статуса
            status_col_idx = 6 + col_idx

            # Строка 1: Количество
            table.rows[row_idx].cells[status_col_idx].text = str(count)

            # Строка 2: Список ФИО (сокращенно)
            if row_idx + 1 < len(table.rows):
                if names:
                    # Сокращаем имена до инициалов
                    short_names = []
                    for name in names[:5]:  # Максимум 5 имен
                        parts = name.split()
                        if len(parts) >= 2:
                            short_name = f"{parts[0]} {parts[1][0]}."
                            if len(parts) > 2:
                                short_name += f"{parts[2][0]}."
                            short_names.append(short_name)
                        else:
                            short_names.append(name)

                    table.rows[row_idx + 1].cells[status_col_idx].text = ', '.join(short_names)
                    if len(names) > 5:
                        table.rows[row_idx + 1].cells[status_col_idx].text += '...'
                else:
                    table.rows[row_idx + 1].cells[status_col_idx].text = '-'

            # Строка 3: Комментарии
            if row_idx + 2 < len(table.rows):
                if comments:
                    comment_text = '; '.join(comments[:2])
                    if len(comments) > 2:
                        comment_text += '...'
                    table.rows[row_idx + 2].cells[status_col_idx].text = comment_text
                else:
                    table.rows[row_idx + 2].cells[status_col_idx].text = '-'

            # Строка 4: Даты
            if row_idx + 3 < len(table.rows):
                if dates:
                    date_text = ', '.join(dates[:3])
                    if len(dates) > 3:
                        date_text += '...'
                    table.rows[row_idx + 3].cells[status_col_idx].text = date_text
                else:
                    table.rows[row_idx + 3].cells[status_col_idx].text = '-'

            # Форматирование всех 4 строк
            for i in range(4):
                if row_idx + i < len(table.rows):
                    cell = table.rows[row_idx + i].cells[status_col_idx]
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)

        row_idx += 4  # Переходим к следующему подразделению

    # Итоговая строка
    if row_idx < len(table.rows):
        total_cells = table.rows[row_idx].cells
        total_cells[0].text = ''
        total_cells[1].text = 'ОБЩЕЕ'
        total_cells[2].text = str(totals['staffing'])
        total_cells[3].text = f"{totals['on_list']} +{totals['seconded_in']}"
        total_cells[4].text = str(totals['vacant'])
        total_cells[5].text = str(totals['in_lineup'])

        for idx, (_, status_type) in enumerate(status_columns):
            total_cells[6 + idx].text = str(totals.get(f'status_{status_type}', 0))

        # Форматирование итоговой строки
        for cell in total_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)


def generate_personnel_report_docx(
    division: Division,
    date_from: datetime.date,
    date_to: Optional[datetime.date] = None
) -> io.BytesIO:
    """
    Генерация отчета расхода личного состава в формате DOCX.

    Если date_to не указан, генерируется отчет на один день.
    Иначе генерируется многостраничный отчет с отдельной страницей на каждый день.
    """
    doc = Document()

    # Установка альбомной ориентации
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    # Настройка полей
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # Определяем диапазон дат
    if date_to is None:
        dates = [date_from]
    else:
        dates = []
        current_date = date_from
        while current_date <= date_to:
            dates.append(current_date)
            current_date += datetime.timedelta(days=1)

    # Генерируем отчет для каждой даты
    for idx, report_date in enumerate(dates):
        if idx > 0:
            # Добавляем разрыв страницы между датами
            doc.add_page_break()

        # Собираем данные по подразделению и его дочерним элементам
        divisions_data = []

        if division.division_type == DivisionType.DEPARTMENT:
            # Для департамента включаем все управления
            managements = division.child_divisions.filter(
                division_type=DivisionType.MANAGEMENT
            ).order_by('name')

            for management in managements:
                divisions_data.append(_collect_division_data(management, report_date))
        else:
            # Для других типов - только само подразделение
            divisions_data.append(_collect_division_data(division, report_date))

        # Создаем таблицу
        _create_report_table_docx(doc, divisions_data, report_date)

    # Сохраняем в буфер
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


def generate_personnel_report_xlsx(
    division: Division,
    date_from: datetime.date,
    date_to: Optional[datetime.date] = None
) -> io.BytesIO:
    """
    Генерация отчета расхода личного состава в формате XLSX.

    Каждая дата на отдельном листе.
    """
    wb = openpyxl.Workbook()
    wb.remove(wb.active)  # Удаляем стандартный лист

    # Определяем диапазон дат
    if date_to is None:
        dates = [date_from]
    else:
        dates = []
        current_date = date_from
        while current_date <= date_to:
            dates.append(current_date)
            current_date += datetime.timedelta(days=1)

    # Стили
    header_font = Font(bold=True, size=14)
    subheader_font = Font(bold=True, size=12)
    data_font = Font(size=10)
    center_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    # Создаем лист для каждой даты
    for report_date in dates:
        ws = wb.create_sheet(title=report_date.strftime('%d.%m.%Y'))

        # Заголовок
        ws.merge_cells('A1:N1')
        title_cell = ws['A1']
        title_cell.value = (
            f"{division.name} ЖЕКЕ ҚҰРАМЫНЫҢ САПТЫҚ ТІЗІМІ "
            f"{report_date.strftime('%d.%m.%Y')} ЖЫЛҒЫ"
        )
        title_cell.font = header_font
        title_cell.alignment = center_alignment

        # Заголовки колонок
        headers = [
            '№', 'Название управления', 'По штату', 'По списку', 'Вакант', 'В строю',
            'На дежурстве', 'После дежурства', 'В командировке',
            'Учёба/соревнования', 'В отпуске', 'На больничном',
            'Прикомандирован', 'Откомандирован'
        ]

        for col, header in enumerate(headers, start=1):
            cell = ws.cell(row=3, column=col)
            cell.value = header
            cell.font = subheader_font
            cell.alignment = center_alignment
            cell.border = border

        # Собираем данные
        divisions_data = []
        if division.division_type == DivisionType.DEPARTMENT:
            managements = division.child_divisions.filter(
                division_type=DivisionType.MANAGEMENT
            ).order_by('name')
            for management in managements:
                divisions_data.append(_collect_division_data(management, report_date))
        else:
            divisions_data.append(_collect_division_data(division, report_date))

        # Заполняем данные
        row = 4
        totals = defaultdict(int)

        for idx, div_data in enumerate(divisions_data, start=1):
            stats = div_data['stats']
            status_details = div_data['status_details']

            # Основные данные
            ws.cell(row=row, column=1).value = idx
            ws.cell(row=row, column=2).value = div_data['division'].name
            ws.cell(row=row, column=3).value = stats['total_staffing']
            ws.cell(row=row, column=4).value = f"{stats['on_list_count']} +{stats['seconded_in_count']}"
            ws.cell(row=row, column=5).value = stats['vacant_count']
            ws.cell(row=row, column=6).value = stats['in_lineup_count']

            # Статусы
            status_columns = [
                EmployeeStatusType.ON_DUTY_ACTUAL,
                EmployeeStatusType.AFTER_DUTY,
                EmployeeStatusType.BUSINESS_TRIP,
                EmployeeStatusType.TRAINING_ETC,
                EmployeeStatusType.ON_LEAVE,
                EmployeeStatusType.SICK_LEAVE,
                EmployeeStatusType.SECONDED_IN,
                EmployeeStatusType.SECONDED_OUT,
            ]

            for col_idx, status_type in enumerate(status_columns, start=7):
                details = status_details.get(status_type, {})
                count = details.get('count', 0)
                if status_type == EmployeeStatusType.SECONDED_IN:
                    count = stats['seconded_in_count']

                ws.cell(row=row, column=col_idx).value = count
                totals[status_type] += count

            # Форматирование строки
            for col in range(1, 15):
                cell = ws.cell(row=row, column=col)
                cell.font = data_font
                cell.alignment = center_alignment
                cell.border = border

            # Обновляем итоги
            totals['staffing'] += stats['total_staffing']
            totals['on_list'] += stats['on_list_count']
            totals['seconded_in'] += stats['seconded_in_count']
            totals['vacant'] += stats['vacant_count']
            totals['in_lineup'] += stats['in_lineup_count']

            row += 1

        # Итоговая строка
        ws.cell(row=row, column=2).value = 'ИТОГО'
        ws.cell(row=row, column=3).value = totals['staffing']
        ws.cell(row=row, column=4).value = f"{totals['on_list']} +{totals['seconded_in']}"
        ws.cell(row=row, column=5).value = totals['vacant']
        ws.cell(row=row, column=6).value = totals['in_lineup']

        for col_idx, status_type in enumerate(status_columns, start=7):
            ws.cell(row=row, column=col_idx).value = totals.get(status_type, 0)

        # Форматирование итоговой строки
        for col in range(1, 15):
            cell = ws.cell(row=row, column=col)
            cell.font = Font(bold=True, size=11)
            cell.alignment = center_alignment
            cell.border = border

        # Настройка ширины колонок
        ws.column_dimensions['A'].width = 5
        ws.column_dimensions['B'].width = 30
        for col in 'CDEFGHIJKLMN':
            ws.column_dimensions[col].width = 12

    # Сохраняем в буфер
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    return buffer


def generate_personnel_report_pdf(
    division: Division,
    date_from: datetime.date,
    date_to: Optional[datetime.date] = None
) -> io.BytesIO:
    """
    Генерация отчета расхода личного состава в формате PDF.
    """
    buffer = io.BytesIO()

    # Настройка документа
    doc = SimpleDocTemplate(
        buffer,
        pagesize=landscape(A4),
        leftMargin=0.5*inch,
        rightMargin=0.5*inch,
        topMargin=0.5*inch,
        bottomMargin=0.5*inch
    )

    # Стили
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        alignment=1,  # Center
        spaceAfter=20
    )

    # Элементы документа
    elements = []

    # Определяем даты
    if date_to is None:
        dates = [date_from]
    else:
        dates = []
        current_date = date_from
        while current_date <= date_to:
            dates.append(current_date)
            current_date += datetime.timedelta(days=1)

    # Генерируем страницы
    for idx, report_date in enumerate(dates):
        if idx > 0:
            elements.append(PageBreak())

        # Заголовок
        title = Paragraph(
            f"{division.name} ЖЕКЕ ҚҰРАМЫНЫҢ САПТЫҚ ТІЗІМІ "
            f"{report_date.strftime('%d.%m.%Y')} ЖЫЛҒЫ",
            title_style
        )
        elements.append(title)

        # Собираем данные
        divisions_data = []
        if division.division_type == DivisionType.DEPARTMENT:
            managements = division.child_divisions.filter(
                division_type=DivisionType.MANAGEMENT
            ).order_by('name')
            for management in managements:
                divisions_data.append(_collect_division_data(management, report_date))
        else:
            divisions_data.append(_collect_division_data(division, report_date))

        # Данные таблицы
        table_data = [[
            '№', 'Управление', 'Штат', 'Список', 'Вакант', 'В строю',
            'Дежурство', 'После деж.', 'Команд.', 'Учёба',
            'Отпуск', 'Больн.', 'Прик.', 'Отком.'
        ]]

        totals = defaultdict(int)

        for idx, div_data in enumerate(divisions_data, start=1):
            stats = div_data['stats']
            status_details = div_data['status_details']

            row = [
                str(idx),
                div_data['division'].name,
                str(stats['total_staffing']),
                f"{stats['on_list_count']}+{stats['seconded_in_count']}",
                str(stats['vacant_count']),
                str(stats['in_lineup_count'])
            ]

            # Статусы
            status_types = [
                EmployeeStatusType.ON_DUTY_ACTUAL,
                EmployeeStatusType.AFTER_DUTY,
                EmployeeStatusType.BUSINESS_TRIP,
                EmployeeStatusType.TRAINING_ETC,
                EmployeeStatusType.ON_LEAVE,
                EmployeeStatusType.SICK_LEAVE,
                EmployeeStatusType.SECONDED_IN,
                EmployeeStatusType.SECONDED_OUT,
            ]

            for status_type in status_types:
                details = status_details.get(status_type, {})
                count = details.get('count', 0)
                if status_type == EmployeeStatusType.SECONDED_IN:
                    count = stats['seconded_in_count']
                row.append(str(count))
                totals[status_type] += count

            table_data.append(row)

            # Обновляем итоги
            totals['staffing'] += stats['total_staffing']
            totals['on_list'] += stats['on_list_count']
            totals['seconded_in'] += stats['seconded_in_count']
            totals['vacant'] += stats['vacant_count']
            totals['in_lineup'] += stats['in_lineup_count']

        # Итоговая строка
        total_row = [
            '', 'ИТОГО',
            str(totals['staffing']),
            f"{totals['on_list']}+{totals['seconded_in']}",
            str(totals['vacant']),
            str(totals['in_lineup'])
        ]
        for status_type in status_types:
            total_row.append(str(totals.get(status_type, 0)))

        table_data.append(total_row)

        # Создаем таблицу
        table = Table(table_data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -2), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
            ('BACKGROUND', (0, -1), (-1, -1), colors.lightgrey),
        ]))

        elements.append(table)

    # Генерируем PDF
    doc.build(elements)
    buffer.seek(0)

    return buffer
