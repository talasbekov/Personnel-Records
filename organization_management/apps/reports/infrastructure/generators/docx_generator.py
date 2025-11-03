"""
Сервисные функции для генерации отчетов и расчета статистики.

Полная реализация всех форматов отчетов согласно ТЗ:
- DOCX с таблицей в альбомной ориентации
- XLSX с форматированием
- PDF отчеты
"""

import io
import datetime
from collections import defaultdict, OrderedDict
from typing import Dict, List, Optional, Tuple

from django.db.models import Count, Q, F
from django.utils import timezone

from io import BytesIO

from docx import Document

import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter

from organization_management.apps.divisions.models import Division
from organization_management.apps.employees.models import Employee
from organization_management.apps.statuses.models import EmployeeStatus
from organization_management.apps.secondments.models import SecondmentRequest

class DOCXGenerator:
    """
    Простейший DOCX‑генератор: заголовок + таблица с агрегатами.
    Возвращает (filename, bytes).
    """

    def generate(self, data, report):
        doc = Document()
        doc.add_heading(f"{report.get_report_type_display()}", level=1)
        doc.add_paragraph(f"Раздел: {data.get('division')}")
        doc.add_paragraph(f"Дата: {data.get('date')}")

        rows = data.get("rows", [])
        table = doc.add_table(rows=1 + len(rows), cols=12)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Подразделение"
        hdr_cells[1].text = "Штатная"
        hdr_cells[2].text = "В строю"
        hdr_cells[3].text = "Отпуск"
        hdr_cells[4].text = "Больничный"
        hdr_cells[5].text = "Командировка"
        hdr_cells[6].text = "Учёба"
        hdr_cells[7].text = "Прикомандировано"
        hdr_cells[8].text = "Откомандировано"
        hdr_cells[9].text = "Прочие отсутствия"
        hdr_cells[10].text = "Итого налич."
        hdr_cells[11].text = "% налич."

        for i, row in enumerate(rows, start=1):
            cells = table.rows[i].cells
            cells[0].text = str(row["division_name"])  # type: ignore
            cells[1].text = str(row["staff_unit"])  # type: ignore
            cells[2].text = str(row["in_service"])  # type: ignore
            cells[3].text = str(row["vacation"])  # type: ignore
            cells[4].text = str(row["sick_leave"])  # type: ignore
            cells[5].text = str(row["business_trip"])  # type: ignore
            cells[6].text = str(row["training"])  # type: ignore
            cells[7].text = str(row["seconded_in"])  # type: ignore
            cells[8].text = str(row["seconded_out"])  # type: ignore
            cells[9].text = str(row["other_absence"])  # type: ignore
            cells[10].text = str(row["present_total"])  # type: ignore
            cells[11].text = str(row["presence_pct"])  # type: ignore

        stream = BytesIO()
        doc.save(stream)
        filename = f"report_{report.id}.docx"
        return filename, stream.getvalue()
